import configparser
import psycopg2
from docx import Document

config = configparser.ConfigParser()
config.read("config.ini")
dbinfo = config["PARSER"]

null_check = lambda x: x if x != '-' else None
is_district = lambda x: any(("федеральный округ" in x, "федерация" in x, "управления" in x, "крымский" not in x))


def parse(file, disease):
    document = Document(file)

    conn = psycopg2.connect(database=dbinfo["database"], user=dbinfo["user"],
                            password=dbinfo["password"], host=dbinfo["host"], port=dbinfo["port"])
    cur = conn.cursor()

    cur.execute("SELECT * FROM federal_districts;")
    districts = {x[0] for x in cur.fetchall() if is_district(x[0])}

    for table in document.tables:
        year1, year2 = 0, 0
        for row in table.rows:
            query1 = list()
            query2 = list()
            t = [cell for cell in iter_cells(row)]
            if "субъекты федерации" in t:
                year1, year2 = null_check(t[1]), null_check(t[2])
                continue
            elif t[0] in districts:
                if disease == "ВИЧ":
                    hiv = True
                else:
                    hiv = False

                if year1:
                    query1 += [disease, year1, t[0]] + list(map(to_num,
                                                                list(map(null_check,
                                                                         t[1:9:2] if hiv else t[1:5:2]))))
                if year2:
                    query2 += [disease, year2, t[0]] + list(map(to_num,
                                                                list(map(null_check,
                                                                         t[2:9:2] if hiv else t[2:5:2]))))

                if query1:
                    execute(conn, cur, query1, hiv)
                if query2:
                    execute(conn, cur, query2, hiv)

    conn.commit()
    cur.close()
    conn.close()


def execute(conn, cur, query, hiv):
    if not hiv:
        try:
            cur.execute(
                "INSERT INTO diseases (disease, year, district, abs, rel, abs_child, rel_child, predicted)"
                " VALUES (%s, %s, %s, %s, %s, NULL, NULL, false);", query)
        except:
            conn.rollback()
            cur.execute("UPDATE diseases"
                        " SET abs = %s, rel = %s, abs_child = %s, rel_child = %s"
                        " WHERE disease = %s AND year = %s AND district = %s",
                        query[:-3:-1] + [None, None] + query[:3])
    else:
        try:
            cur.execute(
                "INSERT INTO diseases (disease, year, district, abs, rel, abs_child, rel_child, predicted)"
                " VALUES (%s, %s, %s, %s, %s, %s, %s, false);", query)
        except:
            conn.rollback()
            cur.execute("UPDATE diseases"
                        " SET abs = %s, rel = %s, abs_child = %s, rel_child = %s"
                        " WHERE disease = %s AND year = %s AND district = %s",
                        query[:-5:-1] + query[:3])


def to_num(string):
    if not string:
        return None
    string = string.replace(',', '.')
    if is_integer_num(string):
        return int(string)
    else:
        return float(string)


def is_integer_num(n):
    if isinstance(n, int):
        return True
    if isinstance(n, float):
        return n.is_integer()
    return False


def iter_cells(row):
    for cell in row.cells:
        yield cell.paragraphs[0].text.strip().lower()
